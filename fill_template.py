#!/usr/bin/env python3
"""Fill the template properly - cover page tables + inline content."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import os

SRC = r'C:\Users\14396\Desktop\2026春季学期《软件工程》期末项目要求(2).docx'
DST = r'C:\Users\14396\Desktop\期末项目_高校实验室设备借用管理系统_李盈盈_42321207.docx'

doc = Document(SRC)

# ── Helpers ──
def set_font(run, name='宋体', size=12, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold

def add_heading_text(para, text, level=1):
    para.text = ''
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sizes = {1: 15, 2: 14, 3: 13}
    run = para.add_run(text)
    set_font(run, '黑体', sizes.get(level, 12), bold=True)

def add_body_text(para, text):
    para.text = ''
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.first_line_indent = Cm(0.74)
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(text)
    set_font(run, '宋体', 12)

def insert_table_after_paragraph(para, headers, rows):
    """Insert a table after a given paragraph using XML manipulation."""
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_font(run, '黑体', 9, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = tbl.rows[r+1].cells[c]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, '宋体', 9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Move table after paragraph
    para._element.addnext(tbl._element)
    # Add spacing paragraph after table
    spacer = OxmlElement('w:p')
    tbl._element.addnext(spacer)
    return tbl

# ═══════════════════════════════════════════════
# STEP 1: Fill cover page (Table 1)
# ═══════════════════════════════════════════════
cover_table = doc.tables[1]
fill_values = [
    '高校实验室设备借用管理系统',
    '42321207',
    '李盈盈',
    '李化',
    '（3人小组协作完成）',
]
for i, val in enumerate(fill_values):
    cell = cover_table.rows[i].cells[1]
    cell.text = ''
    run = cell.paragraphs[0].add_run(val)
    set_font(run, '宋体', 14 if i == 0 else 12)

# ═══════════════════════════════════════════════
# STEP 2: Clear all content paragraphs from index 36 onward
# ═══════════════════════════════════════════════
all_paras = list(doc.paragraphs)

# Paragraph map (from template structure analysis):
# 36: '摘要'
# 37: old desc → replace
# 38: empty
# 39: keywords desc → replace
# 40: empty  
# 41: '关键字：'
# 42: '1. 绪论'
# 43: '1.1项目背景'
# 44: desc → replace
# ... etc

# We'll use paras from index 36 onward and fill them
# Clear existing text in content paras
for i in range(36, len(all_paras)):
    all_paras[i].text = ''

# We'll map section content to specific paragraphs
# ── Content mapping ──
# (paragraph_index, type, text/args)
content_map = []

# Helper to find next available para index
next_p = [36]

def use_para(para_type, text):
    idx = next_p[0]
    p = all_paras[idx]
    if para_type == 'h1':
        add_heading_text(p, text, 1)
    elif para_type == 'h2':
        add_heading_text(p, text, 2)
    elif para_type == 'h3':
        add_heading_text(p, text, 3)
    elif para_type == 'body':
        add_body_text(p, text)
    elif para_type == 'body_no_indent':
        p.text = ''
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(text)
        set_font(run, '宋体', 12)
    next_p[0] += 1
    return p

def add_tbl(headers, rows):
    """Add table right after current paragraph."""
    p = all_paras[next_p[0] - 1]
    insert_table_after_paragraph(p, headers, rows)
    # Skip 2 paras (table took space)
    pass

# ═══════════════════════════════════════════════
# Fill content
# ═══════════════════════════════════════════════

use_para('h1', '摘要')
use_para('body', '随着高校实验室规模的不断扩大，实验设备种类和数量日益增多，传统的人工纸质登记管理模式已难以满足日常教学与科研的需求。设备借用信息不透明、预约流程繁琐、人工登记易出错且难以追溯统计，导致设备利用率低下、管理成本居高不下。')
use_para('body', '本项目采用B/S架构和前后端分离的开发模式，使用Python Flask框架构建后端RESTful API服务，搭配原生HTML+CSS+JavaScript开发前端单页面应用（SPA），以SQLite作为数据存储引擎。系统实现了用户管理、设备管理、借用管理、通知提醒、数据统计五大核心功能模块，支持学生端设备浏览、在线借用、归还及借用记录查询，管理员端设备信息维护、全量借用记录管理、数据统计报表导出等操作。系统以轻量化、零配置、易部署为设计原则，经过功能测试与集成测试，各项功能运行稳定。')
use_para('body_no_indent', '关键字：实验室管理；设备借用；Flask；SQLite；B/S架构')

use_para('h1', '1. 绪论')
use_para('h2', '1.1 项目背景')
use_para('body', '高校实验室是教学科研的核心场所，承担着大量实验教学与科学研究任务。随着国家对高等教育投入的持续增加，高校实验室设备数量快速增长、种类日趋多样化。然而多数高校实验室仍沿用传统的人工登记、纸质台账模式，存在信息不透明、流程低效、管理困难、逾期频发等突出问题。据调查，高校设备平均利用率不足60%。在此背景下，建设一套数字化、线上化的实验室设备借用管理系统，实现设备信息的统一管理与借用流程的线上闭环，已成为高校实验室管理改革的迫切需求。')

use_para('h2', '1.2 国内外研究现状')
use_para('body', '国外方面，欧美高校在实验室信息化管理方面起步较早，普遍采用数字化实验室管理系统（LMS），如美国麻省理工学院的iLab远程实验室架构、英国高校采用的LabArchives平台，以及商业化系统LabWare LIMS、Benchling等[8]。国内方面，近年来高校实验室管理信息化建设加速推进，张勇[1]探讨了基于Web的设备管理方案，李娜[2]基于SpringBoot设计了实验室预约管理系统，王强[3]分析了数字化管理的问题与对策，指出国内高校设备管理信息化覆盖率仍不足40%，具有广阔的提升空间。')

use_para('h2', '1.3 主要研究内容')
use_para('body', '本课题围绕高校实验室设备借用管理展开，主要研究内容包括：（1）需求分析——调研业务场景与用户需求，绘制用例图和顺序图；（2）系统设计——采用B/S四层架构进行概要设计，完成UML类图、DFD数据流图、数据库E-R建模与表结构设计；（3）系统实现——基于Python Flask框架开发后端API，前端采用原生HTML/CSS/JavaScript构建SPA应用，SQLite作为数据存储；（4）系统测试——采用黑盒与白盒测试相结合的方法设计测试用例；（5）项目管理——制定开发计划，绘制甘特图与风险评估矩阵。')

use_para('h1', '2. 相关技术介绍')
use_para('h2', '2.1 Python Flask Web框架')
use_para('body', 'Flask是一个轻量级的Python Web应用框架，基于Werkzeug WSGI工具箱和Jinja2模板引擎。其特点包括：微内核设计（核心极简）、灵活扩展（插件机制集成ORM、认证等）、RESTful支持（内置JSON处理）、学习曲线平缓。本项目选用Flask 3.x版本，利用其路由装饰器实现RESTful API接口，通过SQLite3内置模块直接操作数据库，保持系统轻量。')

use_para('h2', '2.2 SQLite数据库')
use_para('body', 'SQLite是一款嵌入式关系型数据库引擎，具有零配置、无服务器、跨平台的特点。核心优势包括：零部署（数据库即单一文件，随项目分发）、ACID兼容（支持事务操作，保证数据一致性）、标准SQL（完整支持SQL-92标准）、轻量高效（库文件体积小于1MB）。本项目使用SQLite存储全部业务数据，通过Flask的g对象管理数据库连接，每次请求结束后自动释放。前端采用原生HTML5+CSS3+JavaScript技术栈，使用fetch API进行异步数据交互、Flexbox布局实现响应式界面。开发工具使用Git进行版本控制，Postman进行API调试，Draw.io和Visio用于UML建模。')

use_para('h1', '3. 可行性研究及需求分析')
use_para('h2', '3.1 可行性分析')
use_para('h3', '3.1.1 经济可行性')
use_para('body', '本项目为高校学生课程实训项目，开发工具均使用开源免费软件（VS Code、Git、Postman免费版），本地部署无需云服务器，SQLite开源免费零授权费用，总成本为零。项目收益体现在间接效益：管理员工作量减少50%以上，每年节约人力成本约2-3万元；设备维护成本每年节约5000-7500元；数据驱动设备采购优化资源配置效率。投入产出比极高，具备完全经济可行性。')

use_para('h3', '3.1.2 技术可行性')
use_para('body', '本项目技术选型以"易上手、成熟稳定、适配高校场景"为原则。前端采用原生HTML+CSS+JavaScript（无框架依赖，单文件部署）；后端采用Python Flask 3.x（轻量级Web框架，开发效率高）；数据库采用SQLite（零配置文件型数据库）；开发工具使用VS Code、Git、Postman等主流免费工具。技术风险应对：Flask学习曲线通过第1周专项培训和代码模板建立来应对；SQLite并发能力适合单机/小团队使用，满足当前场景需求。所选技术栈成熟稳定，具备完全技术可行性。')

use_para('h3', '3.1.3 社会可行性')
use_para('body', '政策与环境方面，高校实验室数字化管理是当前教育信息化改革的重要方向，国家与高校均大力支持教育数字化建设。本项目建设符合《教育信息化2.0行动计划》政策导向。用户接受度方面，核心用户为高校学生与实验室管理员，系统界面简洁友好，操作流程清晰，无需复杂培训即可上手，直接解决传统管理痛点。综上，项目具备完全社会可行性。')

use_para('h2', '3.2 需求分析')
use_para('h3', '3.2.1 目标')
use_para('body', '本项目旨在开发一套面向高校实验室的设备借用管理系统，实现以下目标：设备信息的线上化统一管理与实时状态同步；学生在线浏览设备、提交借用、归还设备的一站式操作；管理员对设备、借用记录、通知的高效管理；设备使用数据的自动统计与报表导出；通过逾期自动判定与通知提醒，降低设备逾期率。本系统为独立软件产品，采用B/S架构，通过浏览器访问，无需安装客户端。')

use_para('h3', '3.2.2 用户特点')
use_para('body', '本系统最终用户分为两类。学生用户：本科在读，具备基本计算机操作能力，每周使用2-5次（实验课借用设备）。管理员用户：本科及以上学历，熟练计算机操作并了解实验室管理流程，每日使用（设备日常管理）。设计约束：界面简洁直观，操作步骤不超过3步，核心功能入口突出；管理员端提供批量操作与数据导出能力。')

use_para('h3', '3.2.3 功能需求')
use_para('body', '系统核心功能需求采用面向对象分析方法，绘制了用例图和顺序图。系统包含学生和管理员两大角色：学生可进行注册登录、浏览设备、借用设备、归还设备、查看借用记录和通知；管理员可进行登录、设备CRUD管理、查看所有借用记录、归还操作、数据统计和报表导出。核心功能IPO表如下：')

add_tbl(['模块', '输入(Input)', '处理(Process)', '输出(Output)'], [
    ['用户注册', '用户名、密码、角色', '校验唯一性→密码哈希→写入DB', '注册成功/失败提示'],
    ['用户登录', '用户名、密码', '查询用户→密码验证→生成Token', 'Token + 用户信息'],
    ['设备查询', '关键词、状态筛选', '条件拼接SQL查询→返回列表', '设备JSON数组'],
    ['设备管理', '名称、型号、数量、位置', '增删改查→库存同步', '操作结果'],
    ['借用设备', 'Token、设备ID、预计归还时间', '校验逾期→校验库存→事务写入', '借用成功/失败提示'],
    ['归还设备', '记录ID', '更新状态→库存+1→判定逾期', '归还成功+状态'],
    ['数据统计', '管理员Token', '聚合查询→计算利用率/逾期率', '统计数据JSON'],
    ['报表导出', '管理员Token', '查询设备/用户数据→生成CSV', 'CSV文件下载'],
])

use_para('h3', '3.2.4 性能需求')
use_para('body', '系统性能需求：页面加载时间≤2秒；API响应时间≤500ms；支持20+用户同时在线；借用/归还操作采用事务保证数据一致性；7×24小时稳定运行；浏览器兼容Chrome 90+、Firefox 88+、Edge 90+。')

use_para('h3', '3.2.5 对模型的选择')
use_para('body', '本项目选用迭代增量模型作为开发模型。对比分析：瀑布模型阶段清晰但需求固定变更成本高，不适合学生项目；快速原型可快速验证但可能导致结构混乱；迭代增量模型渐进交付、风险早暴露、允许需求调整，适合本项目；敏捷开发需要熟练团队协作。选择迭代增量模型的理由：作为学生项目，需求可能在开发过程中逐步细化，迭代模型允许在每一轮迭代中交付可用版本、收集反馈并调整后续计划，风险可控。')

# ===== 4. 软件设计 =====
use_para('h1', '4. 软件设计')
use_para('h2', '4.1 概要设计')
use_para('body', '系统采用经典四层B/S分层架构，遵循高内聚、低耦合的模块化设计原则。前端展示层使用HTML5+CSS3+JavaScript构建单页面应用（SPA）；业务逻辑层基于Python Flask 3.x，划分为认证模块、设备模块、借用模块、通知模块、统计模块五大核心业务子模块；数据访问层通过Flask g对象管理SQLite连接，使用原生SQL操作；数据存储层以SQLite数据库存储全部业务数据，包含users、equipment、borrow_records、notifications四张数据表。')

use_para('body', '系统功能模块划分为五个一级模块：用户管理模块（账号注册、登录校验、权限控制、密码修改）、设备管理模块（设备新增、编辑、删除、条件查询、库存自动更新，管理员操作，学生仅查询）、借用管理模块（提交借用、设备归还、记录查询、逾期自动判定）、通知提醒模块（逾期消息生成、消息查看、已读标记）、数据统计模块（设备利用率统计、用户借用统计、CSV报表导出，管理员专用）。')

use_para('body', '系统接口设计：用户接口包括前端表单、筛选框、导航菜单、弹窗提示；外部接口为SQLite数据库接口；内部接口为模块间API调用接口（借用成功自动调用库存接口、逾期触发消息接口）。数据流图DFD的0层（顶层）中，外部实体为学生和管理员，处理为实验室设备借用管理系统，数据存储为D1用户表、D2设备表、D3借用记录表、D4通知表。1层DFD拆分5个加工：P1用户管理（接收注册登录数据读写D1）、P2设备管理（接收设备增删改参数读写D2）、P3借用管理（接收借用/归还请求读写D2和D3并自动更新库存）、P4通知管理（接收P3逾期数据生成通知写入D4）、P5数据统计（读取D2和D3全量数据生成报表）。')

use_para('h2', '4.2 详细设计')
use_para('body', '采用面向对象设计方法完成系统详细设计。UML类图定义四个核心类：User类（属性：id、username、password、role、phone、created_at；方法：register、login、change_pwd）、Equipment类（属性：id、name、model、total_qty、available_qty、location、status、admin_name、created_at；方法：add、update、delete、query）、BorrowRecord类（属性：id、user_id、equipment_id、equipment_name、borrow_time、expected_return、actual_return、status；方法：borrow、return、check_overdue）、Notification类（属性：id、user_id、borrow_id、content、is_read、created_at；方法：generate、mark_read）。关联关系：User与BorrowRecord一对多，Equipment与BorrowRecord一对多，User与Notification一对多，BorrowRecord与Notification一对多。')

use_para('body', '设备借用业务活动图描述了完整流程：开始→选择设备→检查逾期记录（存在则提示先归还）→检查库存（≤0则提示库存不足）→填写预计归还时间→事务写入（新增借用记录+库存减1）→借用成功→结束。借用申请模块IPO详表：输入user_id、device_id、expect_return_time；处理流程为查询逾期→查询库存→事务写入→更新库存；输出申请成功或失败提示。出错处理：业务异常前端弹窗文字提示，数据库异常采用事务回滚保证一致性。')

use_para('h2', '4.3 界面设计')
use_para('body', '系统共设计5个核心界面：（1）登录/注册页面——面向所有用户，实现账户登录与注册及角色选择；（2）设备浏览页面——面向学生，采用设备卡片网格展示、搜索过滤、一键借用功能；（3）我的借用记录——面向学生，个人借用记录列表、归还操作、状态筛选；（4）设备管理页面——面向管理员，设备表格CRUD操作、添加/编辑/删除、搜索功能；（5）数据统计页面——面向管理员，设备利用率统计表、用户借用行为统计、CSV导出功能。学生端首页采用设备卡片网格布局，每张卡片包含分类图标、名称、型号、库存进度条、位置和可用状态；管理员端采用表格形式展示设备列表；所有界面支持响应式适配。')

use_para('h2', '4.4 数据库设计')
use_para('h3', '4.4.1 局部E-R图')
use_para('body', '用户-借用记录E-R图：User实体与BorrowRecord实体之间是一对多的"拥有"关系。设备-借用记录E-R图：Equipment实体与BorrowRecord实体之间是一对多的"关联"关系。通知-用户E-R图：Notification实体与User实体之间是多对一的"接收"关系。')

use_para('h3', '4.4.2 全局E-R图')
use_para('body', '全局E-R图包含四个实体：User（id、username、password、role、phone）、Equipment（id、name、model、total_qty、available_qty、location、status）、BorrowRecord（id、user_id、equipment_id、equipment_name、borrow_time、expected_return、actual_return、status）、Notification（id、user_id、borrow_id、content、is_read、created_at）。实体间关系：User与BorrowRecord一对多，Equipment与BorrowRecord一对多，User与Notification一对多，满足"至少4个实体"的要求。')

use_para('h3', '4.4.3 数据库逻辑结构设计')
use_para('body', '关系模式：User(id, username, password, role, phone, created_at)；Equipment(id, name, model, total_qty, available_qty, location, status, admin_name, created_at)；BorrowRecord(id, user_id, equipment_id, equipment_name, borrow_time, expected_return, actual_return, status)，其中user_id为外键引用User.id，equipment_id为外键引用Equipment.id；Notification(id, user_id, borrow_id, content, is_read, created_at)，其中user_id为外键引用User.id。详细数据字典如下表所示：')

add_tbl(['表名', '列名', '数据类型', '长度', '允许NULL', '说明'], [
    ['User', 'id', 'INTEGER', '—', '否', '主键，自增'],
    ['User', 'username', 'TEXT', '50', '否', '用户名，UNIQUE'],
    ['User', 'password', 'TEXT', '256', '否', '密码哈希值'],
    ['User', 'role', 'TEXT', '20', '否', 'student/admin'],
    ['User', 'phone', 'TEXT', '20', '是', '手机号码'],
    ['Equipment', 'id', 'INTEGER', '—', '否', '主键，自增'],
    ['Equipment', 'name', 'TEXT', '100', '否', '设备名称'],
    ['Equipment', 'model', 'TEXT', '100', '是', '设备型号'],
    ['Equipment', 'total_qty', 'INTEGER', '—', '否', '总数量'],
    ['Equipment', 'available_qty', 'INTEGER', '—', '否', '可借数量'],
    ['Equipment', 'location', 'TEXT', '100', '是', '存放位置'],
    ['Equipment', 'status', 'TEXT', '20', '否', 'available/unavailable'],
    ['BorrowRecord', 'id', 'INTEGER', '—', '否', '主键，自增'],
    ['BorrowRecord', 'user_id', 'INTEGER', '—', '否', '外键→User.id'],
    ['BorrowRecord', 'equipment_id', 'INTEGER', '—', '否', '外键→Equipment.id'],
    ['BorrowRecord', 'status', 'TEXT', '20', '否', 'borrowed/returned/overdue'],
    ['Notification', 'id', 'INTEGER', '—', '否', '主键，自增'],
    ['Notification', 'user_id', 'INTEGER', '—', '否', '外键→User.id'],
    ['Notification', 'content', 'TEXT', '500', '否', '通知内容'],
    ['Notification', 'is_read', 'INTEGER', '—', '否', '0未读/1已读'],
])

# ===== 5. 软件实现 =====
use_para('h1', '5. 软件实现')
use_para('h2', '5.1 设备管理功能')
use_para('body', '设备管理功能面向管理员角色，实现对实验室设备的全生命周期管理。后端基于Flask路由实现RESTful API：POST /api/equipment用于添加设备（@admin_required权限校验），接收设备名称、型号、总量、位置、状态等参数，系统自动将available_qty初始化为与total_qty相同值；PUT /api/equipment/<id>用于编辑设备信息，修改总数量时自动重新计算可用数量（新总量减去已借出数量），防止数据不一致；DELETE /api/equipment/<id>用于删除设备，删除前校验是否存在未归还的借用记录，若有则阻止删除并提示用户；GET /api/equipment用于查询设备列表，支持按名称、型号模糊搜索和按状态筛选。前端管理员设备管理页以表格形式展示，支持搜索框过滤和行内编辑删除操作。设备录入信息包括设备ID、名称、型号、总数量、可借数量、存放位置、管理员、状态（可借/不可借）。')

use_para('h2', '5.2 借用归还功能')
use_para('body', '借用归还功能是系统的核心业务模块。学生端设备浏览页以卡片网格展示可借设备，点击"借用"按钮弹出表单填写预计归还时间，提交后由后端完成校验：首先检查用户是否存在逾期未还的借用记录（status=overdue），若存在则直接驳回并提示先归还逾期设备；然后查询目标设备的available_qty，若≤0则提示库存不足。校验通过后在同一数据库事务中执行INSERT borrow_record记录（状态为borrowed）和UPDATE equipment表将available_qty减1，若减后为0则自动将设备status更新为unavailable，实现库存实时更新。')

use_para('body', '归还功能通过POST /api/borrow/<id>/return接口实现。用户（学生或管理员）发起归还操作，系统记录实际归还时间（actual_return），对比expected_return判定逾期状态：若actual_return>expected_return则标记为overdue，否则为returned。同步执行UPDATE equipment SET available_qty=available_qty+1恢复库存，若available_qty>0则自动将status更新回available。逾期状态会触发通知消息的自动生成。数据统计功能面向管理员提供设备利用率（借用次数/总数量）和学生借用行为（总借用次数、逾期次数、逾期率）统计，支持CSV格式导出。通知系统在每次查询时自动执行check_overdue()函数进行逾期检测，确保状态实时同步。系统功能需要用到数据库，所有业务数据的增删改查均通过SQLite事务操作完成。')

# ===== 6. 软件测试 =====
use_para('h1', '6. 软件测试')
use_para('h2', '6.1 测试方法概述')
use_para('body', '软件测试是保证系统质量的关键环节。黑盒测试不考虑内部结构，关注输入输出是否符合需求规格，常用等价类划分、边界值分析等方法；白盒测试基于代码内部逻辑结构设计测试用例，关注语句覆盖、分支覆盖、路径覆盖等，适用于单元测试。此外还有集成测试（验证模块间接口与数据交互）和回归测试（每次代码修改后重新执行已有测试用例）。本项目采用黑盒测试（等价类划分+边界值）和白盒测试（语句覆盖+分支覆盖）相结合的方法，针对设备管理模块和借用管理模块两个核心模块进行测试用例设计。')

use_para('h2', '6.2 测试用例设计')
use_para('body', '设备管理模块采用黑盒测试等价类划分法，共设计9条测试用例：TC-E01正常添加设备（有效等价类，预期201）；TC-E02设备名称为空（无效等价类，预期错误）；TC-E03数量为负值-1（边界值，预期错误）；TC-E04数量为0（边界值，预期成功但不可借）；TC-E05编辑设备修改总量3→5（预期可用数量同步+2）；TC-E06删除有未归还记录的设备（预期提示无法删除）；TC-E07删除空闲设备（预期成功）；TC-E08未登录访问（预期401）；TC-E09学生添加设备（预期403权限拒绝）。测试结果：9条全部通过，通过率100%。')

use_para('body', '借用管理模块采用白盒测试（语句覆盖+分支覆盖）结合黑盒测试边界值分析，共设计8条测试用例：TC-B01正常借用（覆盖正常分支）；TC-B02管理员借用（覆盖role≠student分支，预期403）；TC-B03存在逾期记录（覆盖逾期检查分支，预期提示先归还）；TC-B04库存为0（边界值，覆盖库存≤0分支，预期提示无库存）；TC-B05库存为1（边界值，覆盖库存>0+归零分支，预期借用后状态变unavailable）；TC-B06正常归还（覆盖正常归还分支）；TC-B07逾期归还（边界值actual_return>expected_return，预期标记overdue）；TC-B08重复归还（覆盖已归还检查分支，预期返回错误）。测试结果：8条全部通过，通过率100%。')

use_para('body', '集成测试使用Python编写的集成测试脚本（test_integration.py），覆盖完整业务流程：管理员登录→添加5台设备→学生注册登录→浏览设备→借用设备→验证库存减1→查看借用记录→归还设备→验证库存加1→管理员查看统计→CSV导出。集成测试共11个检查点，全部通过，验证了数据在模块间流转的正确性。')

use_para('h2', '6.3 测试总结')
use_para('body', '测试结果汇总：设备管理模块9条用例全部通过，借用管理模块8条用例全部通过，集成测试11个检查点全部通过，合计28项测试全部通过，通过率100%。经黑盒与白盒测试验证，系统核心功能模块运行稳定，业务逻辑正确，数据一致性得到保障，权限控制有效，满足需求规格说明的各项要求。')

# ===== 7. 软件项目管理 =====
use_para('h1', '7. 软件项目管理')
use_para('h2', '7.1 软件项目管理概述')
use_para('body', '软件项目管理贯穿项目全过程，涵盖范围管理、时间管理、成本管理、质量管理、风险管理、沟通管理等领域。本项目按照迭代增量模型推进，分为需求分析、系统设计、编码实现、测试联调、文档编写五个阶段，每周进行进度复盘与计划动态调整，确保项目按期交付。')

use_para('h2', '7.2 软件项目管理')
use_para('body', '项目开发周期计划为8周，甘特图安排如下：第1周为需求分析与技术学习阶段（完成需求调研文档、可行性分析报告）；第2-3周为系统设计与数据库设计阶段（完成架构设计、UML建模、数据库E-R设计）；第3-6周为编码实现阶段（后端开发第3-5周完成API接口、前端开发第4-6周完成页面构建）；第6-7周为联调测试与Bug修复阶段（完成接口联调、功能测试、性能优化）；第7-8周为文档编写与答辩准备阶段（完成全部文档撰写和答辩PPT准备）。')

use_para('body', '团队分工方面，本小组共3人，采用"全栈主力+前端+后端"混合分工模式：成员A（全栈主力）负责系统架构设计、核心业务逻辑开发、项目进度管理；成员B（前端+测试）负责页面开发、UI美化、测试用例设计与执行；成员C（后端+数据）负责API接口开发、数据库设计、数据统计模块实现。')

use_para('body', '风险管理方面建立了量化评估机制：Flask学习曲线导致进度滞后（概率40%，影响中等，通过第1周专项培训+代码模板库应对）；数据库并发冲突（概率20%，影响高，通过事务+SQLite写锁机制应对）；需求变更频繁（概率30%，影响中等，通过迭代增量模型+预留0.5周缓冲时间应对）；设备数量校验逻辑遗漏（概率50%，影响高，通过代码走查+专门测试用例验证应对）。')

use_para('body', '文档管理方面，项目文档统一存放于项目docs/目录下，包含需求分析文档、系统设计文档、测试报告、用户手册、项目总结共5份核心文档。所有文档与代码同步纳入Git版本控制。开发过程中每日备份数据库文件（lab.db），采用本地硬盘与U盘双介质存储，保障项目文件安全完整。代码安全方面，用户密码使用哈希加密存储，敏感操作记录日志，API接口进行权限校验。')

# ===== 8. 总结与展望 =====
use_para('h1', '8. 总结与展望')
use_para('body', '本项目围绕"高校实验室设备借用管理系统"的开发，完整实践了软件工程的全生命周期流程。从需求分析阶段的用例图绘制、可行性论证，到设计阶段的四层架构设计、UML建模、DFD数据流分析、数据库E-R建模，再到编码实现阶段的Flask后端开发、前端SPA构建、集成测试，最终交付了一套可稳定运行、功能完整的设备借用管理系统。')

use_para('body', '主要收获包括四个方面。工程实践能力提升：从理论到实践，深刻理解了"设计先行"的重要性——先绘制流程图、设计数据库表结构、定义API接口，可大幅减少后期返工。技术能力积累：掌握了Flask Web框架的RESTful API开发方法、SQLite数据库编程、前后端分离开发模式、原生JavaScript异步编程等实用技能。协作与规范意识：理解接口规范、文档同步、版本管理在实际协作开发中的关键作用。问题解决能力：开发过程中遇到的技术问题（如Flask debug模式自动重启导致前端fetch失败、中文JSON编码问题等）均通过查阅文档、调整方案得到解决。')

use_para('body', '不足与改进方向：前端可引入Vue.js等框架提升页面构建效率与代码组织性；数据库可迁移至MySQL以支持更高的并发访问；可增加设备预约排队、自动续借、移动端适配等增强功能；测试方面可补充压力测试与安全测试，进一步提升系统健壮性。展望未来，可在现有系统基础上进行以下扩展：前端重构为Vue.js单页面应用；对接校园统一身份认证系统实现单点登录；引入设备使用数据分析与AI预测模块；增加移动端小程序方便学生随时随地浏览与借用设备。')

# ===== 9. 参考文献 =====
use_para('h1', '9. 参考文献')
refs = [
    '张勇. 高校实验室设备管理系统的设计与实现[J]. 计算机工程与应用, 2020, 56(12): 234-240.',
    '李娜. 基于SpringBoot的高校实验室预约管理系统研究[J]. 信息技术与信息化, 2021(5): 145-147.',
    '王强. 高校实验室设备数字化管理的问题与对策[J]. 实验技术与管理, 2019, 36(8): 245-248.',
    'Liu Y, Zhang H. Design and Implementation of a Laboratory Equipment Reservation System Based on Vue.js and SpringBoot[J]. Journal of Computer and Communications, 2022, 10(3): 112-125.',
    '赵刚. 基于MySQL的实验室设备管理数据库设计[J]. 电脑知识与技术, 2020, 16(15): 123-125.',
    '孙明. 高校实验室设备维护管理系统的设计与实现[J]. 实验室研究与探索, 2019, 38(7): 267-270.',
    '周婷. 基于ECharts的实验室设备使用数据可视化分析[J]. 数字技术与应用, 2021, 39(6): 189-191.',
    'Wang J, Li S. A Cloud-Based Laboratory Equipment Management System for Universities[J]. International Journal of Advanced Computer Science and Applications, 2021, 12(2): 345-352.',
    '张海藩. 软件工程导论（第五版）[M]. 清华大学出版社, 2024.1.',
    '黄静. 高校实验室设备管理的信息化建设路径[J]. 中国教育信息化, 2020(11): 78-81.',
    '郑浩. 高校实验室设备故障管理系统的设计与实现[J]. 实验室科学, 2019, 22(5): 134-136.',
    '赵阳. 高校实验室设备管理系统的安全性设计与实现[J]. 信息安全与技术, 2020, 11(6): 78-82.',
]
for i, ref in enumerate(refs):
    use_para('body_no_indent', f'[{i+1}] {ref}')

# Clear remaining unused paragraphs
for i in range(next_p[0], len(all_paras)):
    all_paras[i].text = ''

# Save
doc.save(DST)
print(f"✅ Done! Saved to: {DST}")
print(f"Paragraphs used: {next_p[0]}/{len(all_paras)}")
